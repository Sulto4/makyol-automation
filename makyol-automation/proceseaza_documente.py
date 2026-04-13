#!/usr/bin/env python3
"""
Automatizare Makyol - Procesare documente materiale pentru Supervizor Tehnic
============================================================================
Citeste PDF-uri din subfoldere (un folder per furnizor), extrage informatii
cu pdfplumber + OCR fallback si genereaza Word cu adresa + tabel centralizator.

Utilizare:
    python proceseaza_documente.py --input "Documente" --output "Adresa_Materiale_Supervizor.docx"

Dependinte:
    pip install pdfplumber python-docx pytesseract pdf2image Pillow
"""

import os
import re
import argparse
import json
from datetime import datetime
from PIL import Image

# Increase PIL limit for large scans
Image.MAX_IMAGE_PIXELS = 300000000

try:
    import pdfplumber
except ImportError:
    print("EROARE: pip install pdfplumber")
    exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("EROARE: pip install python-docx")
    exit(1)

# OCR imports (optional but recommended)
HAS_OCR = False
try:
    import pytesseract
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    print("AVERTISMENT: OCR indisponibil. Instaleaza: pip install pytesseract pdf2image")
    print("  Documentele scanate vor fi clasificate doar dupa numele fisierului.\n")


# ============================================================================
# CONFIGURARE OCR
# ============================================================================

def setup_ocr():
    """Configureaza calea Tesseract si verifica disponibilitatea."""
    if not HAS_OCR:
        return False

    # Cauta tessdata in locatii cunoscute
    tessdata_paths = [
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tesseract-ocr/5/tessdata',
        '/usr/local/share/tessdata',
    ]

    # Verifica si custom tessdata
    custom_tessdata = os.environ.get('TESSDATA_PREFIX', '')
    if custom_tessdata:
        tessdata_paths.insert(0, custom_tessdata)

    for path in tessdata_paths:
        eng_path = os.path.join(path, 'eng.traineddata')
        if os.path.exists(eng_path):
            os.environ['TESSDATA_PREFIX'] = path
            return True

    return False


def ocr_pdf(filepath, max_pages=5, dpi=200):
    """Extrage text din PDF scanat folosind OCR."""
    if not HAS_OCR:
        return ""

    try:
        images = convert_from_path(filepath, dpi=dpi, first_page=1, last_page=max_pages)
        all_text = ""
        # Foloseste ron+eng daca e disponibil, altfel doar eng
        lang = 'eng'
        tessdata = os.environ.get('TESSDATA_PREFIX', '')
        if tessdata and os.path.exists(os.path.join(tessdata, 'ron.traineddata')):
            lang = 'ron+eng'

        for img in images:
            text = pytesseract.image_to_string(img, lang=lang)
            all_text += text + "\n"
        return all_text
    except Exception as e:
        print(f"    OCR EROARE: {e}")
        return ""


# ============================================================================
# EXTRAGERE TEXT DIN PDF
# ============================================================================

def extract_text(filepath, max_pages=5):
    """Extrage text din PDF: mai intai pdfplumber, apoi OCR fallback."""
    text = ""

    # Pasul 1: pdfplumber
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages[:max_pages]:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception:
        pass

    # Pasul 2: OCR fallback daca textul e insuficient
    if len(text.strip()) < 50:
        ocr_text = ocr_pdf(filepath)
        if ocr_text:
            text = ocr_text

    return text


# ============================================================================
# CLASIFICARE DOCUMENTE
# ============================================================================

FILENAME_PATTERNS = [
    (r'(?i)(?:\bFT\b.*\bPEHD\b|\bFT\b.*\bteav|fisa\s*tehnica|\bFT\b.*\bfiting|data\s*sheet)', 'Fisa tehnica'),
    (r'(?i)ISO\s*9001', 'Certificat ISO 9001'),
    (r'(?i)ISO\s*14001', 'Certificat ISO 14001'),
    (r'(?i)ISO\s*45001', 'Certificat ISO 45001'),
    (r'(?i)ISO\s*50001', 'Certificat ISO 50001'),
    (r'(?i)(?:\bAVS\b|aviz\s*sanitar)', 'Aviz sanitar'),
    (r'(?i)(?:\bAVT\b|aviz\s*tehnic)', 'Aviz tehnic'),
    (r'(?i)(?:\bAGT\b|agrement\s*tehnic|\bAT\b\s*\d{3}|aviz.*agrement)', 'Agrement tehnic'),
    (r'(?i)(?:\bDC\b(?!\s*-)|declarati[ea]\s*de\s*conformitate)', 'Declaratie de conformitate'),
    (r'(?i)(?:\bCC\b(?:\s*-|\.)|certificat\s*de\s*calitate)', 'Certificat de calitate'),
    (r'(?i)(?:\bCG\b|certificat\s*de\s*garantie)', 'Certificat de garantie'),
    (r'(?i)(?:certificat.*P1R|certificat.*CERT)', 'Certificat de conformitate'),
    (r'(?i)(?:aprobare.*teav|aprobare.*PEID|aprobare.*produc)', 'Aprobare material'),
    (r'(?i)autorizare\s*de\s*comercializare', 'Autorizare comercializare'),
]

TEXT_PATTERNS = [
    (r'(?i)autorizare\s*de\s*comercializare', 'Autorizare comercializare'),
    (r'(?i)(?:fi[sș][aă]\s*tehnic[aă]|data\s*sheet|technical\s*data)', 'Fisa tehnica'),
    (r'(?i)ISO\s*9001', 'Certificat ISO 9001'),
    (r'(?i)ISO\s*14001', 'Certificat ISO 14001'),
    (r'(?i)ISO\s*45001', 'Certificat ISO 45001'),
    (r'(?i)ISO\s*50001', 'Certificat ISO 50001'),
    (r'(?i)aviz\s*sanitar', 'Aviz sanitar'),
    (r'(?i)aviz\s*tehnic', 'Aviz tehnic'),
    (r'(?i)agrement\s*tehnic', 'Agrement tehnic'),
    (r'(?i)declara[tț]i[ea]\s*de\s*conformitate', 'Declaratie de conformitate'),
    (r'(?i)certificat\s*de\s*(?:conformitate|calitate)', 'Certificat de calitate'),
    (r'(?i)certificat\s*de\s*garan[tț]ie', 'Certificat de garantie'),
    (r'(?i)supervizorul\s*aprob[aă]', 'Aprobare material'),
]


def classify_document(filename, text=""):
    """Clasificare document pe baza numelui si continutului."""
    # Incearca mai intai dupa filename
    for pattern, doc_type in FILENAME_PATTERNS:
        if re.search(pattern, filename):
            return doc_type

    # Apoi pe baza textului
    if text:
        for pattern, doc_type in TEXT_PATTERNS:
            if re.search(pattern, text):
                return doc_type

    return "Document tehnic"


# ============================================================================
# EXTRAGERE DATE STRUCTURATE
# ============================================================================

MONTH_MAP_RO = {
    'ianuarie': '01', 'februarie': '02', 'martie': '03', 'aprilie': '04',
    'mai': '05', 'iunie': '06', 'iulie': '07', 'august': '08',
    'septembrie': '09', 'octombrie': '10', 'noiembrie': '11', 'decembrie': '12',
}


def normalize_date(date_str):
    """Normalizeaza o data in format DD.MM.YYYY."""
    if not date_str:
        return ""

    date_str = date_str.strip()

    # YYYY-MM-DD (cu posibile spatii)
    m = re.match(r'(\d{4})\s*[-./]\s*(\d{1,2})\s*[-./]\s*(\d{1,2})', date_str)
    if m:
        return f"{m.group(3).zfill(2)}.{m.group(2).zfill(2)}.{m.group(1)}"

    # DD.MM.YYYY or DD/MM/YYYY
    m = re.match(r'(\d{1,2})\s*[./]\s*(\d{1,2})\s*[./]\s*(\d{4})', date_str)
    if m:
        return f"{m.group(1).zfill(2)}.{m.group(2).zfill(2)}.{m.group(3)}"

    # DD luna YYYY (romanian)
    m = re.match(r'(\d{1,2})\s+(\w+)\s+(\d{4})', date_str)
    if m:
        month = MONTH_MAP_RO.get(m.group(2).lower(), '')
        if month:
            return f"{m.group(1).zfill(2)}.{month}.{m.group(3)}"

    return date_str


def extract_expiry_date(text):
    """Extrage data de expirare din text cu pattern-uri multiple."""
    if not text:
        return ""

    patterns = [
        # Romanian: "valabil/a pana la (data de) DD.MM.YYYY"
        r'(?i)valabil[aă]?\s*(?:p[aâ]n[aă])\s*(?:la|[iî]n)\s*(?:data\s*(?:de)?\s*)?(\d{1,2}\s*[./]\s*\d{1,2}\s*[./]\s*\d{4})',
        # Romanian: "pana la data de DD luna YYYY"
        r'(?i)p[aâ]n[aă]\s*la\s*(?:data\s*(?:de)?\s*)?(\d{1,2}\s+\w+\s+\d{4})',
        # Romanian: "valabil din X pana in DD.MM.YYYY"
        r'(?i)valabil\s*(?:din\s*[\d./ ]+)?\s*p[aâ]n[aă]\s*[iî]n\s*(\d{1,2}\s*[./]\s*\d{1,2}\s*[./]\s*\d{4})',
        # English: "valid until YYYY-MM-DD"
        r'(?i)(?:valid\s*until|expires?\s*on)[:\s]*(\d{4}\s*[-./]\s*\d{1,2}\s*[-./]\s*\d{1,2})',
        # English: "Expires on: DD month YYYY"
        r'(?i)expires?\s*on[:\s]*(\d{1,2}\s+\w+\s+\d{4})',
        # Certificat pattern: "valabil pana la DD.MM.YYYY"
        r'(?i)valabil[aă]?\s*p[aâ]n[aă]\s*la[:\s]*(\d{1,2}\s*[./]\s*\d{1,2}\s*[./]\s*\d{4})',
        # "ramane valabil pana la data de DD.MM.YYYY"
        r'(?i)r[aă]m[aâ]ne\s*valabil\s*p[aâ]n[aă]\s*la\s*(?:data\s*(?:de)?\s*)?(\d{1,2}\s*[./]\s*\d{1,2}\s*[./]\s*\d{4})',
        # Fallback: "pana la DD.MM.YYYY" near "valabil"
        r'(?i)p[aâ]n[aă]\s*la\s*(\d{1,2}\s*[./]\s*\d{1,2}\s*[./]\s*\d{4})',
    ]

    # Collect all matches, take the LAST one (usually the most specific / latest expiry)
    all_matches = []
    for pat in patterns:
        matches = re.findall(pat, text)
        all_matches.extend(matches)

    if all_matches:
        # Normalize and pick the latest date
        dates = []
        for m in all_matches:
            normalized = normalize_date(m)
            if normalized and re.match(r'\d{2}\.\d{2}\.\d{4}', normalized):
                dates.append(normalized)

        if dates:
            # Sort and return latest
            try:
                sorted_dates = sorted(dates, key=lambda d: datetime.strptime(d, '%d.%m.%Y'), reverse=True)
                return sorted_dates[0]
            except ValueError:
                return dates[0]

    return ""


def extract_material_from_text(text, doc_type, folder_name):
    """Extrage materialul specific din text."""
    if text:
        # Pattern-uri specifice pentru materiale
        specific_patterns = [
            # Tevi POLITUB multistrat
            r'(?i)(TEVI\s*"?POLITUB"?\s*MULTISTRAT\s*DIN\s*POLIETILEN[AĂ][^\n]{0,30})',
            # Tevi WaterKIT / VALWater
            r'(?i)((?:Tevi|Țevi)\s*(?:<|")?(?:WaterKIT|VALWater)(?:>|")?\s*[^\n]{0,50})',
            # Tevi si fitinguri VALROM
            r'(?i)((?:Tevi|Țevi)\s*[sș]i\s*fitinguri\s*VALROM\s*[^\n]{0,60})',
            # Tevi din PEID/PEHD generic
            r'(?i)((?:Tevi|Țevi)\s*(?:din\s*)?PE(?:ID|HD)\s*(?:PE\s*100\s*(?:RC|CERT)?)?(?:\s*(?:pentru|pt\.?)\s*(?:instalati[ei]|ap[aă]))?[^\n]{0,30})',
            # Tevi monostrat/multistrat
            r'(?i)((?:Tevi|Țevi)\s*(?:monostrat|multistrat)\s*[^\n]{0,60})',
            # Fitinguri segmentate
            r'(?i)(FITINGURI\s*SEGMENTATE[^\n]{0,40})',
            # Pipes - english
            r'(?i)(PE100\s*(?:RC|CERT)?\s*(?:VALWater|WaterKIT)?\s*PIPES[^\n]{0,40})',
        ]

        for pat in specific_patterns:
            match = re.search(pat, text)
            if match:
                result = match.group(1).strip()
                # Curata si trunchiaza
                result = re.sub(r'\s+', ' ', result)
                return result[:80]

    # Fallback pe baza folderului
    FOLDER_MATERIALS = {
        'TERAPLAST': 'Tevi PEHD PE100 pentru apa',
        'VALROM': 'Tevi PEHD PE100/PE100RC pentru apa',
        'TEHNOWORLD': 'Tevi PEHD PE100 pentru apa',
        'TEHNO WORLD': 'Tevi PEHD PE100 pentru apa',
        'ZAKPREST': 'Tevi PEHD PE100 pentru apa',
    }

    folder_upper = folder_name.upper()
    for key, value in FOLDER_MATERIALS.items():
        if key in folder_upper:
            return value

    return ""


def extract_producer(text, folder_name):
    """Extrage producatorul din text sau folder."""
    if text:
        patterns = [
            r'(?i)produc[aă]tor[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
            r'(?i)titularul\s*certificatului[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
            r'(?i)fabricant[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
            r'(?i)fabricat[ea]?\s*de\s*(?:catre\s*)?(?:firma\s*)?(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
            r'(?i)Solicitant[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
        ]
        for pat in patterns:
            match = re.search(pat, text)
            if match:
                result = match.group(1).strip()
                # Filtrare false positives - max 60 chars, fara cuvinte irelevante
                if (len(result) > 5 and len(result) < 60 and
                    not any(w in result for w in ['AVIZ', 'CERTIFICAT', 'MINISTERUL',
                                                   'utilizarea', 'Etichetarea', 'conformitate'])):
                    return result

    KNOWN_PRODUCERS = {
        'TERAPLAST': 'TERAPLAST SA',
        'VALROM': 'VALROM INDUSTRIE SRL',
        'TEHNOWORLD': 'TEHNO WORLD SRL',
        'TEHNO WORLD': 'TEHNO WORLD SRL',
        'ZAKPREST': 'TERAPLAST SA',  # Zakprest e distribuitor Teraplast
    }

    folder_upper = folder_name.upper()
    for key, value in KNOWN_PRODUCERS.items():
        if key in folder_upper:
            return value

    return folder_name


def extract_distributor(text, folder_name):
    """Extrage distribuitorul din text sau context."""
    # Zakprest e distribuitor autorizat Teraplast
    if 'ZAKPREST' in folder_name.upper():
        return 'ZAKPREST CONSTRUCT SRL'

    # Cauta in text
    if text:
        patterns = [
            r'(?i)distribuit(?:or)?[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
            r'(?i)furnizor[:\s]+(?:S\.?C\.?\s*)?([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
        ]
        for pat in patterns:
            match = re.search(pat, text)
            if match:
                result = match.group(1).strip()
                if len(result) > 5 and 'AVIZ' not in result:
                    return result

    # Producatori directi - fara intermediar
    return "-"


def detect_authorization_docs(results):
    """Detecteaza relatii producator-distribuitor din documentele de autorizare."""
    authorizations = {}

    for info in results:
        if info['doc_type'] == 'Autorizare comercializare':
            # Cauta relatia in text
            text = info.get('_raw_text', '')
            # Pattern: "autorizeaza ... Societatea X ... sa livreze"
            match = re.search(
                r'(?i)autorizeaz[aă]\s+.*?Societatea\s+([A-Z][A-Z\s.]+(?:S\.?R\.?L\.?|S\.?A\.?))',
                text
            )
            if match:
                distributor = match.group(1).strip()
                producer = info.get('producer', '')
                authorizations[info['folder']] = {
                    'producer': producer,
                    'distributor': distributor,
                }

    return authorizations


# ============================================================================
# PROCESARE
# ============================================================================

def process_pdf(filepath, folder_name):
    """Proceseaza un singur PDF si extrage informatiile."""
    filename = os.path.basename(filepath)

    # Extrage text (pdfplumber + OCR fallback)
    text = extract_text(filepath)

    # Clasificare
    doc_type = classify_document(filename, text)

    # Extrage date
    expiry = extract_expiry_date(text)
    producer = extract_producer(text, folder_name)
    distributor = extract_distributor(text, folder_name)
    material = extract_material_from_text(text, doc_type, folder_name)

    return {
        'filename': filename,
        'folder': folder_name,
        'doc_type': doc_type,
        'material': material,
        'expiry': expiry if expiry else '-',
        'producer': producer,
        'distributor': distributor,
        'has_text': len(text.strip()) > 50,
        '_raw_text': text,  # pastram pentru detectia autorizarilor
    }


def process_folder(input_dir):
    """Proceseaza toate PDF-urile din toate subfolderele."""
    results = []

    for folder_name in sorted(os.listdir(input_dir)):
        folder_path = os.path.join(input_dir, folder_name)
        if not os.path.isdir(folder_path) or folder_name.startswith('.'):
            continue

        print(f"\nProcesez: {folder_name}")
        print("-" * 60)

        for filename in sorted(os.listdir(folder_path)):
            if not filename.lower().endswith('.pdf'):
                continue

            filepath = os.path.join(folder_path, filename)
            print(f"  -> {filename}", flush=True)

            info = process_pdf(filepath, folder_name)
            results.append(info)

            status = []
            status.append(f"Tip: {info['doc_type']}")
            if info['expiry'] != '-':
                status.append(f"Expira: {info['expiry']}")
            if info['producer']:
                status.append(f"Prod: {info['producer']}")
            if info['distributor'] != '-':
                status.append(f"Dist: {info['distributor']}")
            if info['has_text']:
                status.append("(text)")
            else:
                status.append("(OCR)")
            print(f"     {' | '.join(status)}")

    return results


# ============================================================================
# GENERARE WORD
# ============================================================================

def generate_word(results, output_path, project_info=None):
    """Genereaza documentul Word cu adresa oficiala si tabelul."""

    if project_info is None:
        project_info = {
            'nr_adresa': 'MAKYOL-RO-LOT4-UTI-___',
            'data': datetime.now().strftime('%d.%m.%Y'),
            'proiect': 'Proiectare si Executie Autostrada Sibiu – Fagaras',
            'tronson': 'Tronsonul 4: Sambata de Sus (DJ 105B) – Municipiul Fagaras km 51+785 – km 68+050',
            'destinatar': 'TECNIC Consulting Engineering Romania SRL',
            'subiect': 'Transmitere documente materiale pentru relocare conducte apa',
        }

    doc = Document()

    # Setari pagina (A4 landscape for wider table)
    section = doc.sections[0]
    section.page_height = Cm(21)
    section.page_width = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Stil default
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- ANTET ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('MAKYOL INSAAT SANAYI TURIZM VE TICARET A.S.')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('YAPI MERKENZI INSAAT VE SANAYI A.S.')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # Nr si data
    p = doc.add_paragraph()
    run = p.add_run(f"Nr. {project_info['nr_adresa']}")
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run(f"Data: {project_info['data']}")
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # Destinatar
    p = doc.add_paragraph()
    run = p.add_run(f"Catre: {project_info['destinatar']}")
    run.bold = True
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('In atentia: Supervizor Tehnic')
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # Referinta proiect
    p = doc.add_paragraph()
    run = p.add_run(f"Ref.: {project_info['proiect']}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    run = p.add_run(f"{project_info['tronson']}")
    run.font.name = 'Arial'
    run.font.size = Pt(10)

    doc.add_paragraph('')

    # Subiect
    p = doc.add_paragraph()
    run = p.add_run(f"Subiect: {project_info['subiect']}")
    run.bold = True
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # Corp adresa
    p = doc.add_paragraph()
    run = p.add_run('Stimati Domni,')
    run.font.name = 'Arial'

    doc.add_paragraph('')

    p = doc.add_paragraph()
    text_body = (
        'Va transmitem anexat documentele tehnice ale materialelor care vor fi '
        'utilizate pentru relocarea conductelor de apa in cadrul proiectului mentionat mai sus. '
        'Documentele includ fise tehnice, certificate ISO, agremente tehnice, avize sanitare, '
        'declaratii de conformitate si certificate de garantie pentru materialele propuse.'
    )
    run = p.add_run(text_body)
    run.font.name = 'Arial'

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Lista documentelor transmise este urmatoarea:')
    run.font.name = 'Arial'

    doc.add_paragraph('')

    # --- TABEL ---
    headers = ['Nr.', 'Nume document', 'Material', 'Data expirare\n(daca e cazul)',
               'Producator', 'Distribuitor']

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E75B6',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    # Separator rows per folder
    current_folder = None
    row_idx = 0

    for info in results:
        # Adauga separator de folder
        if info['folder'] != current_folder:
            current_folder = info['folder']
            sep_row = table.add_row()
            # Merge cells for separator
            sep_cell = sep_row.cells[0]
            sep_cell.merge(sep_row.cells[5])
            sep_cell.text = ''
            p = sep_cell.paragraphs[0]
            run = p.add_run(f"  {current_folder}")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(44, 62, 80)

            shading = sep_cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): 'D6E4F0',
                qn('w:val'): 'clear',
            })
            shading.append(shading_elem)

        row_idx += 1
        row = table.add_row()
        cells = row.cells

        data = [
            str(row_idx),
            f"{info['doc_type']}\n({info['filename']})",
            info['material'],
            info['expiry'],
            info['producer'],
            info['distributor'],
        ]

        for i, val in enumerate(data):
            cell = cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Zebra striping
            if row_idx % 2 == 0:
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F7F9FC',
                    qn('w:val'): 'clear',
                })
                shading.append(shading_elem)

    # Latimi coloane (landscape A4: 29.7 - 4 = 25.7 cm content)
    widths = [Cm(1.2), Cm(5.5), Cm(5), Cm(3), Cm(5.5), Cm(5.5)]
    for row in table.rows:
        for i, width in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = width

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Incheiere
    p = doc.add_paragraph()
    run = p.add_run(
        'Va rugam sa analizati documentele anexate si sa ne transmiteti aprobarea '
        'pentru utilizarea materialelor sus-mentionate in cadrul lucrarilor de '
        'relocare conducte apa.'
    )
    run.font.name = 'Arial'

    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Cu stima,')
    run.font.name = 'Arial'

    doc.add_paragraph('')
    doc.add_paragraph('')

    p = doc.add_paragraph()
    run = p.add_run('Reprezentant Antreprenor')
    run.bold = True
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    run = p.add_run('Asocierea MAKYOL INSAAT - YAPI MERKENZI')
    run.font.name = 'Arial'

    doc.save(output_path)
    print(f"\nDocument salvat: {output_path}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Proceseaza documente PDF materiale si genereaza Word pentru Supervizor Tehnic'
    )
    parser.add_argument('--input', '-i', default='Documente',
                        help='Folderul cu subfoldere de documente (default: Documente)')
    parser.add_argument('--output', '-o', default='Adresa_Materiale_Supervizor.docx',
                        help='Fisierul Word de output (default: Adresa_Materiale_Supervizor.docx)')
    parser.add_argument('--json', '-j',
                        help='Exporta si un JSON cu datele extrase (optional)')
    parser.add_argument('--no-ocr', action='store_true',
                        help='Dezactiveaza OCR (clasificare doar pe baza numelui fisierului)')

    args = parser.parse_args()

    input_dir = args.input
    if not os.path.isdir(input_dir):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        input_dir = os.path.join(script_dir, args.input)
        if not os.path.isdir(input_dir):
            print(f"EROARE: Folderul '{args.input}' nu exista!")
            exit(1)

    print("=" * 60)
    print("AUTOMATIZARE MAKYOL - Procesare documente materiale")
    print("=" * 60)
    print(f"Input:  {input_dir}")
    print(f"Output: {args.output}")

    # Setup OCR
    if args.no_ocr:
        global HAS_OCR
        HAS_OCR = False
        print("OCR:    DEZACTIVAT")
    elif HAS_OCR:
        if setup_ocr():
            print("OCR:    ACTIV (Tesseract)")
        else:
            print("OCR:    Tesseract gasit dar tessdata lipseste")
    else:
        print("OCR:    INDISPONIBIL")

    # Procesare
    results = process_folder(input_dir)

    # Curata raw text din rezultate inainte de export
    export_results = []
    for r in results:
        clean = {k: v for k, v in r.items() if k != '_raw_text'}
        export_results.append(clean)

    print(f"\n{'=' * 60}")
    print(f"Total documente procesate: {len(results)}")
    print(f"  - Cu text extractibil:   {sum(1 for r in results if r['has_text'])}")
    print(f"  - Fara text (clasificate dupa nume): {sum(1 for r in results if not r['has_text'])}")

    # Generare Word
    output_path = args.output
    if not os.path.isabs(output_path):
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_path = os.path.join(script_dir, output_path)

    generate_word(results, output_path)

    # Optional: export JSON
    if args.json:
        json_path = args.json
        if not os.path.isabs(json_path):
            script_dir = os.path.dirname(os.path.abspath(__file__))
            json_path = os.path.join(script_dir, json_path)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(export_results, f, ensure_ascii=False, indent=2)
        print(f"JSON exportat: {json_path}")


if __name__ == '__main__':
    main()
