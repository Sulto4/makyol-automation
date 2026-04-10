"""Word document generation for the PDF processing automation pipeline.

Generates a formatted .docx output file with a cover letter and summary table
listing all processed construction material documents.
"""

import logging
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.models import DocumentInfo

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

HEADER_BG_COLOR = "4472C4"
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_COLOR = "D9E2F3"
FONT_NAME = "Arial"

TABLE_HEADERS = [
    "Nr.",
    "Nume document",
    "Material",
    "Data expirare\n(daca e cazul)",
    "Producator",
    "Distribuitor",
]

# Default project info
DEFAULT_PROJECT_INFO = {
    "company_1": "MAKYOL INSAAT SANAYI TURIZM VE TICARET A.S.",
    "company_2": "YAPI MERKENZI INSAAT VE SANAYI A.S.",
    "reference": "YM-MKL/SF T4/",
    "addressee": "TECNIC Consulting Engineering Romania SRL",
    "attention": "Supervizor Tehnic",
    "project": (
        "Proiectare si Executie Autostrada Sibiu-Fagaras, "
        "Tronsonul 4"
    ),
    "subject": (
        "Transmitere documente materiale pentru "
        "relocare conducte apa"
    ),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _set_cell_shading(cell, color_hex: str) -> None:
    """Set the background shading of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def _set_cell_text(cell, text: str, bold: bool = False,
                   font_size: int = 10, color: RGBColor | None = None,
                   alignment: int | None = None) -> None:
    """Set cell text with formatting."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_styled_paragraph(doc: Document, text: str, bold: bool = False,
                          font_size: int = 11,
                          alignment: int | None = None,
                          space_after: int | None = None) -> None:
    """Add a paragraph with consistent styling."""
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    if space_after is not None:
        paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold


# ---------------------------------------------------------------------------
# Main generation function
# ---------------------------------------------------------------------------


def generate_word_output(
    documents: list[DocumentInfo],
    output_path: Path,
    project_info: dict | None = None,
) -> None:
    """Generate the output Word document with cover letter and summary table.

    Args:
        documents: List of DocumentInfo objects to include in the table.
        output_path: Path where the .docx file will be saved.
        project_info: Optional dict overriding default project metadata.

    Raises:
        OSError: If the output file cannot be written.
    """
    info = {**DEFAULT_PROJECT_INFO, **(project_info or {})}
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Set default font for the document --
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    # ----- Cover letter / header section -----

    # Company names
    _add_styled_paragraph(
        doc, info["company_1"], bold=True, font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
    )
    _add_styled_paragraph(
        doc, info["company_2"], bold=True, font_size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )

    # Reference and date
    _add_styled_paragraph(
        doc, f"Ref: {info['reference']}", font_size=11, space_after=2,
    )
    _add_styled_paragraph(doc, "Data: _______________", font_size=11, space_after=12)

    # Addressee
    _add_styled_paragraph(doc, "Catre:", bold=True, font_size=11, space_after=2)
    _add_styled_paragraph(doc, info["addressee"], font_size=11, space_after=2)
    _add_styled_paragraph(
        doc, f"In atentia: {info['attention']}", font_size=11, space_after=12,
    )

    # Project reference
    _add_styled_paragraph(doc, "Referinta proiect:", bold=True, font_size=11, space_after=2)
    _add_styled_paragraph(doc, info["project"], font_size=11, space_after=12)

    # Subject
    _add_styled_paragraph(
        doc, f"Subiect: {info['subject']}", bold=True, font_size=11, space_after=12,
    )

    # Body text
    _add_styled_paragraph(
        doc,
        "Va transmitem alaturat documentele aferente materialelor "
        "utilizate pentru relocarea conductelor de apa, conform tabelului de mai jos:",
        font_size=11,
        space_after=12,
    )

    # ----- Summary table -----

    num_cols = len(TABLE_HEADERS)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_row = table.rows[0]
    for idx, header_text in enumerate(TABLE_HEADERS):
        cell = header_row.cells[idx]
        _set_cell_shading(cell, HEADER_BG_COLOR)
        _set_cell_text(
            cell, header_text, bold=True, font_size=10,
            color=HEADER_TEXT_COLOR,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Data rows
    for row_num, doc_info in enumerate(documents, start=1):
        row = table.add_row()
        values = [
            str(row_num),
            doc_info.file_name,
            doc_info.material or "N/A",
            doc_info.expiration_date or "N/A",
            doc_info.producer or "N/A",
            doc_info.distributor or "N/A",
        ]

        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            _set_cell_text(cell, value, font_size=10)

            # Zebra striping on even data rows (0-indexed: row_num 2,4,6...)
            if row_num % 2 == 0:
                _set_cell_shading(cell, ZEBRA_COLOR)

    # ----- Closing / signature block -----

    _add_styled_paragraph(doc, "", space_after=12)  # spacer
    _add_styled_paragraph(
        doc, "Cu stima,", font_size=11, space_after=24,
    )
    _add_styled_paragraph(
        doc, info["company_1"], bold=True, font_size=11, space_after=2,
    )
    _add_styled_paragraph(
        doc, info["company_2"], bold=True, font_size=11, space_after=24,
    )
    _add_styled_paragraph(doc, "Nume: _______________", font_size=11, space_after=2)
    _add_styled_paragraph(doc, "Semnatura: _______________", font_size=11, space_after=2)
    _add_styled_paragraph(doc, "Data: _______________", font_size=11)

    # Save
    doc.save(str(output_path))
    logger.info("Word document saved to %s (%d documents)", output_path, len(documents))
