"""Unit tests for word_generator module."""

import pytest
from pathlib import Path

from src.word_generator import generate_word_output, TABLE_HEADERS
from src.models import DocumentInfo


class TestGenerateWordOutput:
    """Test Word document generation."""

    def _make_doc_info(self, **kwargs):
        defaults = {
            "file_path": "/fake/path.pdf",
            "file_name": "test_doc.pdf",
            "supplier_folder": "TERAPLAST",
            "material": "Tevi PEHD PE100",
            "expiration_date": "31.12.2025",
            "producer": "TERAPLAST SA",
            "distributor": "N/A",
        }
        defaults.update(kwargs)
        return DocumentInfo(**defaults)

    def test_creates_output_file(self, tmp_path):
        """Word file is created at the specified path."""
        output = tmp_path / "output" / "result.docx"
        docs = [self._make_doc_info()]
        generate_word_output(docs, output)
        assert output.exists()

    def test_creates_parent_directories(self, tmp_path):
        """Output parent directories are created automatically."""
        output = tmp_path / "deep" / "nested" / "result.docx"
        generate_word_output([], output)
        assert output.exists()

    def test_table_has_correct_column_count(self, tmp_path):
        """Table should have 6 columns matching TABLE_HEADERS."""
        from docx import Document

        output = tmp_path / "result.docx"
        docs = [self._make_doc_info()]
        generate_word_output(docs, output)

        doc = Document(str(output))
        tables = doc.tables
        assert len(tables) == 1
        table = tables[0]
        assert len(table.columns) == 6
        assert len(table.columns) == len(TABLE_HEADERS)

    def test_table_header_row_text(self, tmp_path):
        """Header row should contain the expected column names."""
        from docx import Document

        output = tmp_path / "result.docx"
        generate_word_output([self._make_doc_info()], output)

        doc = Document(str(output))
        header_row = doc.tables[0].rows[0]
        header_texts = [cell.text.strip() for cell in header_row.cells]
        expected = [h.replace("\n", " ").strip() for h in TABLE_HEADERS]
        # Compare without newline differences
        for actual, exp in zip(header_texts, expected):
            assert exp.split()[0] in actual

    def test_data_row_values(self, tmp_path):
        """Data rows should contain document info values."""
        from docx import Document

        output = tmp_path / "result.docx"
        doc_info = self._make_doc_info(
            file_name="ISO 9001.pdf",
            material="Tevi PEHD",
            producer="TERAPLAST SA",
        )
        generate_word_output([doc_info], output)

        doc = Document(str(output))
        data_row = doc.tables[0].rows[1]
        cells = [cell.text for cell in data_row.cells]
        assert cells[0] == "1"  # Nr.
        assert cells[1] == "ISO 9001.pdf"  # Nume document
        assert cells[2] == "Tevi PEHD"  # Material
        assert cells[4] == "TERAPLAST SA"  # Producator

    def test_multiple_documents(self, tmp_path):
        """Multiple documents create multiple data rows."""
        from docx import Document

        output = tmp_path / "result.docx"
        docs = [self._make_doc_info(file_name=f"doc{i}.pdf") for i in range(5)]
        generate_word_output(docs, output)

        doc = Document(str(output))
        # 1 header + 5 data rows
        assert len(doc.tables[0].rows) == 6

    def test_empty_documents_list(self, tmp_path):
        """Empty document list creates file with header-only table."""
        from docx import Document

        output = tmp_path / "result.docx"
        generate_word_output([], output)

        doc = Document(str(output))
        assert len(doc.tables[0].rows) == 1  # header only
